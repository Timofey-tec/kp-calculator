"""Генерация договора: заполнение contract_template.docx через python-docx.

Никакого шаблонизатора (jinja/docxtpl) не используется — плейсхолдеры вида
{{KEY}} заменяются вручную по тексту параграфов, а таблица сметы
вставляется в OOXML-дерево на место параграфа-маркера {{SMETA_TABLE}}.
"""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt

from app.schemas import ClientInfo
from app.services.calculations import QuoteCalculation
from app.services.number_to_words import amount_to_words_ru

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "templates" / "contract_template.docx"

TABLE_HEADERS = ["Наименование", "Кол-во", "Цена за ед.", "Скидка, %", "Сумма"]


def format_money(value: float) -> str:
    return f"{value:,.2f}".replace(",", " ").replace(".", ",")


def _replace_placeholders_in_paragraph(paragraph, mapping: dict[str, str]) -> None:
    if not any(f"{{{{{key}}}}}" in paragraph.text for key in mapping):
        return
    new_text = paragraph.text
    for key, value in mapping.items():
        new_text = new_text.replace(f"{{{{{key}}}}}", value)

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return

    first_run = paragraph.runs[0]
    first_run.text = new_text
    for extra_run in paragraph.runs[1:]:
        extra_run.text = ""


def _replace_placeholders(document: Document, mapping: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        _replace_placeholders_in_paragraph(paragraph, mapping)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _replace_placeholders_in_paragraph(paragraph, mapping)


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.size = Pt(10)


def _insert_smeta_table(document: Document, calculation: QuoteCalculation, vat_enabled: bool, vat_rate: float) -> None:
    placeholder = None
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == "{{SMETA_TABLE}}":
            placeholder = paragraph
            break
    if placeholder is None:
        raise ValueError("Плейсхолдер {{SMETA_TABLE}} не найден в шаблоне договора")

    extra_rows = 1  # "Итого"
    if not calculation.delivery_distributed and calculation.delivery_cost_total > 0:
        extra_rows += 1
    if vat_enabled:
        extra_rows += 1
    extra_rows += 1  # "Итого с НДС" / итоговая сумма к оплате

    table = document.add_table(rows=1 + len(calculation.items) + extra_rows, cols=5)
    table.style = "Table Grid"

    for col, header in enumerate(TABLE_HEADERS):
        _set_cell_text(table.rows[0].cells[col], header, bold=True)

    row_idx = 1
    for item in calculation.items:
        cells = table.rows[row_idx].cells
        _set_cell_text(cells[0], item.name)
        _set_cell_text(cells[1], f"{item.quantity:g}")
        _set_cell_text(cells[2], format_money(item.unit_price))
        _set_cell_text(cells[3], f"{item.discount_percent:g}")
        _set_cell_text(cells[4], format_money(item.final_line_total))
        row_idx += 1

    def add_summary_row(label: str, value: str, bold: bool = False) -> None:
        nonlocal row_idx
        cells = table.rows[row_idx].cells
        cells[0].merge(cells[3])
        _set_cell_text(cells[0], label, bold=bold)
        _set_cell_text(cells[4], value, bold=bold)
        row_idx += 1

    add_summary_row("Итого", format_money(calculation.subtotal_after_delivery - (
        calculation.delivery_cost_total if not calculation.delivery_distributed else 0
    )))

    if not calculation.delivery_distributed and calculation.delivery_cost_total > 0:
        add_summary_row("Доставка", format_money(calculation.delivery_cost_total))

    if vat_enabled:
        add_summary_row(f"НДС ({vat_rate:g}%)", format_money(calculation.vat_amount))

    add_summary_row("Итого к оплате", format_money(calculation.grand_total), bold=True)

    # переносим только что созданную (добавленную в конец документа) таблицу
    # на место параграфа-маркера, затем сам параграф удаляем
    placeholder._p.addnext(table._tbl)
    placeholder._p.getparent().remove(placeholder._p)


def generate_contract(
    client: ClientInfo,
    calculation: QuoteCalculation,
    vat_enabled: bool,
    vat_rate: float,
    contract_number: str = "1",
) -> io.BytesIO:
    document = Document(TEMPLATE_PATH)

    mapping = {
        "CONTRACT_NUMBER": contract_number,
        "DATE": date.today().strftime("%d.%m.%Y"),
        "COMPANY_NAME": client.company_name,
        "INN": client.inn or "не указан",
        "CONTACT_PERSON": client.contact_person,
        "PHONE": client.phone or "не указан",
        "EMAIL": client.email or "не указан",
        "TOTAL_DIGITS": format_money(calculation.grand_total),
        "TOTAL_WORDS": amount_to_words_ru(calculation.grand_total),
    }

    _insert_smeta_table(document, calculation, vat_enabled, vat_rate)
    _replace_placeholders(document, mapping)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer
