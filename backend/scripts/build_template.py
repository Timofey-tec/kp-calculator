"""Одноразовый скрипт: собирает backend/app/templates/contract_template.docx.

Запускается вручную (`python scripts/build_template.py`) при необходимости
поправить текст шаблона договора. Результат коммитится в репозиторий, чтобы
демо работало сразу после клонирования без ручной подготовки шаблона.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

TEMPLATE_PATH = Path(__file__).resolve().parent.parent / "app" / "templates" / "contract_template.docx"

SUPPLIER_NAME = 'ООО «ПромТехСнаб»'


def add_heading(doc: Document, text: str, size: int = 14, bold: bool = True, center: bool = False) -> None:
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)


def build() -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    add_heading(doc, "ДОГОВОР ПОСТАВКИ № {{CONTRACT_NUMBER}}", size=16, center=True)
    add_body(doc, "г. Москва\t\t\t\t\t\t\t\t{{DATE}}")
    doc.add_paragraph()

    add_body(
        doc,
        f"{SUPPLIER_NAME}, именуемое в дальнейшем «Поставщик», с одной стороны, и "
        "{{COMPANY_NAME}}, ИНН {{INN}}, именуемое в дальнейшем «Покупатель», в лице "
        "{{CONTACT_PERSON}}, с другой стороны, совместно именуемые «Стороны», "
        "заключили настоящий Договор о нижеследующем.",
    )
    doc.add_paragraph()

    add_heading(doc, "1. ПРЕДМЕТ ДОГОВОРА", size=12)
    add_body(
        doc,
        "1.1. Поставщик обязуется передать в собственность Покупателя товар в "
        "соответствии со спецификацией (сметой), являющейся неотъемлемой частью "
        "настоящего Договора (Приложение № 1), а Покупатель обязуется принять "
        "товар и оплатить его на условиях настоящего Договора.",
    )
    doc.add_paragraph()

    add_heading(doc, "2. СТОИМОСТЬ И ПОРЯДОК РАСЧЁТОВ", size=12)
    add_body(
        doc,
        "2.1. Стоимость товара по настоящему Договору определяется спецификацией "
        "(сметой), приведённой ниже.",
    )
    doc.add_paragraph("{{SMETA_TABLE}}")
    add_body(doc, "2.2. Итого к оплате: {{TOTAL_DIGITS}} руб. ({{TOTAL_WORDS}}).")
    add_body(
        doc,
        "2.3. Оплата производится Покупателем в течение 5 (пяти) рабочих дней с "
        "момента подписания настоящего Договора путём безналичного перечисления "
        "денежных средств на расчётный счёт Поставщика.",
    )
    doc.add_paragraph()

    add_heading(doc, "3. КОНТАКТНЫЕ ДАННЫЕ ПОКУПАТЕЛЯ", size=12)
    add_body(doc, "Контактное лицо: {{CONTACT_PERSON}}")
    add_body(doc, "Телефон: {{PHONE}}")
    add_body(doc, "Email: {{EMAIL}}")
    doc.add_paragraph()

    add_heading(doc, "4. ЗАКЛЮЧИТЕЛЬНЫЕ ПОЛОЖЕНИЯ", size=12)
    add_body(
        doc,
        "4.1. Настоящий Договор вступает в силу с момента подписания Сторонами и "
        "действует до полного исполнения Сторонами своих обязательств.",
    )
    add_body(
        doc,
        "4.2. Все изменения и дополнения к настоящему Договору действительны "
        "лишь в том случае, если они совершены в письменной форме и подписаны "
        "надлежаще уполномоченными представителями Сторон.",
    )
    doc.add_paragraph()
    doc.add_paragraph()

    add_heading(doc, "5. ПОДПИСИ СТОРОН", size=12)
    doc.add_paragraph()
    add_body(doc, f"Поставщик: {SUPPLIER_NAME}\t\t\tПокупатель: {{{{COMPANY_NAME}}}}")
    doc.add_paragraph()
    add_body(doc, "_____________________ / _____________ /\t\t_____________________ / {{CONTACT_PERSON}} /")

    TEMPLATE_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(TEMPLATE_PATH)
    print(f"Шаблон сохранён: {TEMPLATE_PATH}")


if __name__ == "__main__":
    build()
